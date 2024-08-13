import streamlit as st
import pandas as pd
import time
import google.generativeai as genai
from io import BytesIO
from docx import Document
from fpdf import FPDF

# Configurer la page de l'application
st.set_page_config(layout="wide")

# Fonction pour ajouter des styles CSS personnalisés
def add_css_styles():
    st.markdown(
        """
        <style>
        .main-header {
            font-size: 24px;
            font-weight: bold;
            color: #004080;
            text-align: center;
            margin-bottom: 25px;
        }
        .banner {
            text-align: center;
            margin-bottom: 40px;
        }
        .dataframe-container {
            margin-bottom: 40px;
        }
        .recommendation-container {
            padding: 20px;
            margin-bottom: 30px;
            border: 1px solid #004080;
            border-radius: 5px;
            background-color: #f0f8ff;
        }
        .recommendation-header {
            font-weight: bold;
            font-size: 18px;
            color: #004080;
            margin-bottom: 10px;
        }
        .recommendation-content {
            margin-bottom: 10px;
            font-size: 16px;
            line-height: 1.5;
        }
        .spinner-message {
            font-size: 22px;
            color: red;
            text-align: center;
            margin-top: 20px;
        }
        .warning {
            color: red;
            font-weight: bold;
        }
        .success {
            color: green;
            font-weight: bold;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

# Fonction pour configurer le modèle d'IA de Google
def configure_model():
    genai.configure(api_key=st.secrets["api_key"])
    model = genai.GenerativeModel(model_name="gemini-1.5-pro-latest")
    return model

# Fonction pour charger le fichier Excel avec le plan d'action
def load_action_plan(uploaded_file):
    try:
        action_plan_df = pd.read_excel(uploaded_file, header=12)
        expected_columns = ["Numéro d'exigence", "Exigence IFS Food 8", "Notation", "Explication (par l’auditeur/l’évaluateur)"]
        action_plan_df = action_plan_df[expected_columns]
        return action_plan_df
    except Exception as e:
        st.error(f"Erreur lors de la lecture du fichier: {str(e)}")
        return None

# Fonction pour générer un prompt basé sur une non-conformité
def generate_prompt(non_conformity, guide_df):
    exigence_num = non_conformity.get("Numéro d'exigence", "Exigence non spécifiée")
    non_conformity_desc = non_conformity.get("Exigence IFS Food 8", "Non-conformité non spécifiée")
    
    context = guide_df[guide_df['NUM_REQ'].str.contains(exigence_num, na=False)]
    guidelines = context['Good practice'].iloc[0] if not context.empty else "Aucune information spécifique trouvée dans le guide pour cette exigence."

    prompt = f"""
    Je suis un expert en IFS Food 8. Voici une non-conformité détectée lors d'un audit :
    - **Exigence** : {exigence_num}
    - **Non-conformité** : {non_conformity_desc}
    Contexte supplémentaire tiré du guide IFSv8 :
    {guidelines}
    Fournissez une recommandation structurée et détaillée comprenant :
    - **Correction immédiate** (action spécifique et claire)
    - **Preuves requises** (documents précis nécessaires)
    - **Actions Correctives** (mesures à long terme)
    """
    
    return prompt

# Fonction pour générer une recommandation à partir de l'IA
def generate_ai_recommendation(prompt, model):
    try:
        convo = model.start_chat(history=[{"role": "user", "parts": [prompt]}])
        response = convo.send_message(prompt)
        return response.text if response else None
    except Exception as e:
        st.error(f"Erreur lors de la génération de la recommandation: {str(e)}")
        return None

# Fonction pour afficher les recommandations avec un rendu Markdown
def display_recommendations(recommendations_df):
    for index, row in recommendations_df.iterrows():
        st.markdown(f"""### Numéro d'exigence: {row["Numéro d'exigence"]}""")
        st.markdown(row["Recommandation"])

# Fonction pour créer un fichier texte des recommandations
def generate_text_file(recommendations_df):
    text_content = ""
    for index, row in recommendations_df.iterrows():
        text_content += f"Numéro d'exigence: {row['Numéro d'exigence']}\n"
        text_content += f"{row['Recommandation']}\n\n"
    return text_content

# Fonction pour créer un fichier DOCX des recommandations
def generate_docx_file(recommendations_df):
    doc = Document()
    for index, row in recommendations_df.iterrows():
        doc.add_heading(f"Numéro d'exigence: {row['Numéro d'exigence']}", level=2)
        doc.add_paragraph(row['Recommandation'])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Fonction pour créer un fichier PDF des recommandations
def generate_pdf_file(recommendations_df):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for index, row in recommendations_df.iterrows():
        pdf.set_font("Arial", style='B', size=14)
        pdf.cell(200, 10, txt=f"Numéro d'exigence: {row['Numéro d'exigence']}", ln=True)
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=row['Recommandation'])
        pdf.ln(10)
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer

# Fonction principale
def main():
    add_css_styles()

    st.markdown('<div class="banner"><img src="https://raw.githubusercontent.com/M00N69/Gemini-Knowledge/main/visipilot%20banner.PNG" alt="Banner" width="80%"></div>', unsafe_allow_html=True)
    st.markdown('<div class="main-header">Assistant VisiPilot pour Plan d\'Actions IFS</div>', unsafe_allow_html=True)
    st.write("Téléchargez votre plan d'action et obtenez des recommandations pour les corrections et les actions correctives.")

    uploaded_file = st.file_uploader("Téléchargez votre plan d'action (fichier Excel)", type=["xlsx"])
    
    if uploaded_file:
        action_plan_df = load_action_plan(uploaded_file)
        if action_plan_df is not None:
            st.markdown('<div class="dataframe-container">' + action_plan_df.to_html(classes='dataframe', index=False) + '</div>', unsafe_allow_html=True)
            model = configure_model()

            # Charger le guide IFSv8 depuis le CSV sur GitHub
            guide_df = pd.read_csv("https://raw.githubusercontent.com/M00N69/Action-plan/main/Guide%20Checklist_IFS%20Food%20V%208%20-%20CHECKLIST.csv")

            # Préparation d'une liste pour les recommandations
            recommendations = []

            # Affichage d'un spinner et d'un message d'attente pendant la génération des recommandations
            with st.spinner('Génération des recommandations en cours...'):
                st.markdown('<p class="spinner-message">Veuillez patienter pendant que les recommandations sont générées...</p>', unsafe_allow_html=True)
                for _, non_conformity in action_plan_df.iterrows():
                    prompt = generate_prompt(non_conformity, guide_df)
                    recommendation_text = generate_ai_recommendation(prompt, model)
                    if recommendation_text:
                        recommendations.append({
                            "Numéro d'exigence": non_conformity["Numéro d'exigence"],
                            "Recommandation": recommendation_text
                        })
                    # Pause entre les requêtes pour éviter l'épuisement des ressources
                    time.sleep(5)  # Attendre 5 secondes entre chaque requête

            if recommendations:
                recommendations_df = pd.DataFrame(recommendations)
                st.subheader("Résumé des Recommandations")
                display_recommendations(recommendations_df)

                # Télécharger au format CSV
                st.download_button(
                    label="Télécharger les recommandations (CSV)",
                    data=recommendations_df.to_csv(index=False),
                    file_name="recommandations_ifs_food.csv",
                    mime="text/csv",
                )

                # Télécharger au format texte
                text_file = generate_text_file(recommendations_df)
                st.download_button(
                    label="Télécharger les recommandations (Texte)",
                    data=text_file,
                    file_name="recommandations_ifs_food.txt",
                    mime="text/plain",
                )

                # Télécharger au format DOCX
                docx_file = generate_docx_file(recommendations_df)
                st.download_button(
                    label="Télécharger les recommandations (DOCX)",
                    data=docx_file,
                    file_name="recommandations_ifs_food.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

                # Télécharger au format PDF
                pdf_file = generate_pdf_file(recommendations_df)
                st.download_button(
                    label="Télécharger les recommandations (PDF)",
                    data=pdf_file,
                    file_name="recommandations_ifs_food.pdf",
                    mime="application/pdf",
                )

if __name__ == "__main__":
    main()







